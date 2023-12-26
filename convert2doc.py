import requests
from xml.etree import ElementTree as ET
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def parse_time(time_str):
    """Convert time string to datetime object."""
    try:
        return datetime.strptime(time_str, "%Y%m%d%H%M%S %z")
    except ValueError:
        # Handle invalid or empty time string
        return None

def calculate_duration(start, end):
    """Calculate the duration between two datetime objects in hours."""
    if start and end:
        duration_seconds = (end - start).total_seconds()
        return duration_seconds / 3600  # Convert seconds to hours
    else:
        return 0

def format_datetime(dt):
    """Format datetime object into a human readable form."""
    if dt:
        return dt.strftime("%Y-%m-%d %I:%M %p")
    else:
        return "Unknown"

def fetch_and_parse_xml(url):
    """Fetch and parse XML data from the given URL."""
    response = requests.get(url)
    return ET.fromstring(response.content)

def create_word_document(root, file_path):
    """Create a Word document with programme details."""
    doc = Document()
    doc.add_heading('SCBC Television Programme Schedule', 0)

    for programme in root.iter('programme'):
        start_time = parse_time(programme.get('start'))
        end_time = parse_time(programme.get('stop'))

        # Check if the start time is within the desired range (6 AM to 11:59:59 PM)
        if not start_time or not end_time or not (6 <= start_time.hour < 24):
            continue

        # Adjust end time if it goes past midnight
        if end_time.hour < 6:
            end_time = end_time.replace(hour=23, minute=59, second=59)

        title = programme.find('title').text
        duration = calculate_duration(start_time, end_time)

        p = doc.add_paragraph()
        p.add_run(f"Title: {title}\n").bold = True
        p.add_run(f"Start Time: {format_datetime(start_time)}\n")
        p.add_run(f"End Time: {format_datetime(end_time)}\n")
        p.add_run(f"Duration: {duration:.2f} hours\n")
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    doc.save(file_path)

# URL of the XMLTV DTD data
url = "https://raw.githubusercontent.com/digitalp/scbcguide/main/scbcguide.xml"

# Path where the Word document will be saved
output_file_path = "C:\\epg\\schedule.docx"

# Fetch, parse, and save the programme schedule
xml_root = fetch_and_parse_xml(url)
create_word_document(xml_root, output_file_path)

print(f"Schedule saved to {output_file_path}")
